"""
ResumeBuilderAgent — creates tailored Word + PDF resumes.

Takes the CompanyAgent's tailoring brief and builds a polished DOCX.
Inspired by career-ops pdf/latex modes + the docx skill.
"""

from __future__ import annotations
import os
import logging
from datetime import date
from pathlib import Path
from typing import Optional
import yaml

from agents.base_agent import BaseAgent
from config.settings import get_settings

logger = logging.getLogger(__name__)


class ResumeBuilderAgent(BaseAgent):
    """Builds tailored Word resumes from CompanyAgent briefs."""

    def __init__(self):
        settings = get_settings()
        super().__init__(name="ResumeBuilder", model=settings.rizwan_agent_model)
        self.profile = self._load_profile()

    def _load_profile(self) -> dict:
        try:
            with open(get_settings().profile_path) as f:
                return yaml.safe_load(f)
        except FileNotFoundError:
            return {}

    async def run(
        self,
        job_title: str,
        company: str,
        tailoring_brief: str,
        job_id: Optional[int] = None
    ) -> str:
        """Build the tailored resume and return file path."""
        self.log(f"Building resume for {job_title} @ {company}")

        # Parse the tailoring brief to extract structured data
        sections = await self._parse_tailoring_brief(tailoring_brief, job_title, company)

        # Build the DOCX
        path = self._build_docx(sections, job_title, company)
        self.log(f"✅ Resume saved: {path}", style="green")

        # Workflow v2: upload to Supabase Storage so it survives Railway redeploys
        try:
            from db.client import upload_artifact
            remote_path = f"resumes/{Path(path).name}"
            url = upload_artifact(
                local_path=path,
                remote_path=remote_path,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            if url:
                self.log(f"  📤 Uploaded to: {url[:80]}", style="cyan")
                # Return the URL — the caller will store it in resume_path
                return url
        except Exception as e:
            self.log(f"  ⚠️ Upload failed: {e}", style="yellow")

        return path

    async def _parse_tailoring_brief(self, brief: str, job_title: str, company: str) -> dict:
        """Use Claude to extract structured resume sections from the brief."""
        system = "You are a resume structuring assistant. Extract clean, formatted resume sections from a brief."
        user = f"""Extract structured resume data from this tailoring brief for {job_title} at {company}:

{brief}

Return JSON with these keys:
{{
  "summary": "3-4 sentence professional summary",
  "top_skills": ["skill1", "skill2", ...],  // max 12
  "featured_bullets": [  // top 5-7 most relevant experience bullets
    {{"company": "...", "achievement": "bullet text with metric"}}
  ],
  "new_bullets": [  // new bullets from gap-filling
    {{"company": "SimPaisa", "achievement": "..."}}
  ],
  "cover_hook": "one specific sentence for the cover email"
}}"""

        try:
            response = await self.ask_claude(
                system=system,
                messages=[{"role": "user", "content": user}],
                max_tokens=2000,
                temperature=0.2,
            )
            import json
            start = response.find("{")
            end = response.rfind("}") + 1
            return json.loads(response[start:end])
        except Exception as e:
            logger.warning(f"Brief parsing failed: {e}. Using defaults.")
            return {}

    def _build_docx(self, sections: dict, job_title: str, company: str) -> str:
        """Build the DOCX file using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return ""

        p = self.profile.get("candidate", {})
        exp = self.profile.get("experience", {})
        certs = self.profile.get("certifications", [])

        doc = Document()

        # Page margins
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

        # ── Header ──────────────────────────────────────────────────────────
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(p.get("full_name", "Rizwan Zafar").upper())
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact = f"{p.get('phone', '')} • {p.get('email', '')} • {p.get('location', '')} • {p.get('linkedin', '')}"
        contact_run = contact_para.add_run(contact)
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # ── Section divider function ─────────────────────────────────────────
        def add_section_header(title: str):
            para = doc.add_paragraph()
            run = para.add_run(title.upper())
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1A1A2E")
            pBdr.append(bottom)
            pPr.append(pBdr)
            para.space_after = Pt(4)
            return para

        # ── Professional Summary ─────────────────────────────────────────────
        add_section_header("Professional Summary")
        summary_text = sections.get("summary") or (
            f"Results-driven Technical Programme Lead & CPO with {exp.get('total_years', 14)}+ years "
            f"building and scaling digital payments platforms across CEMEA. Currently leading SimPaisa "
            f"to $1B+ TPV with Fortune 500 integrations (TikTok, Uber, MoneyGram). Expert in product "
            f"strategy, programme governance, and enterprise-grade delivery. PMP, PMI-ACP, CSPO, CSM certified."
        )
        para = doc.add_paragraph(summary_text)
        para.style.font.size = Pt(10)
        para.space_after = Pt(4)

        # ── Core Competencies ────────────────────────────────────────────────
        add_section_header("Core Competencies")
        top_skills = sections.get("top_skills") or [
            "Product Strategy & Roadmapping", "Programme Governance (RAID/SteerCo)",
            "Digital Payments & Fintech", "Agile / Scrum / SAFe",
            "Stakeholder Management", "Enterprise Integration Delivery",
            "Team Leadership (40+ engineers)", "KYC/AML/PCI-DSS Compliance",
            "A/B Testing & Data Analytics", "OKR Frameworks", "MS Project / Think-cell"
        ]
        # 3-column table for skills
        skill_table = doc.add_table(rows=1, cols=3)
        skill_table.style = "Table Grid"
        rows_needed = (len(top_skills) + 2) // 3
        for i in range(rows_needed):
            row = skill_table.add_row() if i > 0 else skill_table.rows[0]
            for col_idx in range(3):
                skill_idx = i * 3 + col_idx
                if skill_idx < len(top_skills):
                    cell = row.cells[col_idx]
                    cell.text = f"✓ {top_skills[skill_idx]}"
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
        doc.add_paragraph()  # spacing

        # ── Professional Experience ──────────────────────────────────────────
        add_section_header("Professional Experience")

        # Full experience sections
        experiences = [
            {
                "title": "Technical Programme Lead / Chief Product Officer",
                "company": "SimPaisa",
                "location": "Karachi, Pakistan",
                "period": "August 2020 – Present",
                "bullets": [
                    "Scaled digital payments platform from 0 to $1B+ total payment volume (TPV) in 4 years",
                    "Led 40+ engineers across 12 cross-functional squads (mobile, backend, payments, risk, compliance)",
                    "Delivered 4 Fortune 500 enterprise integrations: TikTok, Uber, MoneyGram, PUBG — on time, within SLA",
                    "Owned full programme governance: RAID log across 50+ dependencies, SteerCo ownership, monthly board reporting",
                    "Built SBP-compliant payment infrastructure: 3DS, PCI-DSS, KYC/AML workflows",
                    "Launched BNPL product from 0 to 100K users in 8 months through agile product discovery and rapid iteration",
                    "Ran OKR framework, product roadmapping, feature prioritisation (RICE/MoSCoW), and quarterly planning",
                ],
            },
            {
                "title": "Senior Product Manager",
                "company": "Daraz / Alibaba Group",
                "location": "Karachi, Pakistan",
                "period": "2018 – 2020",
                "bullets": [
                    "Led checkout, payments, and seller portal product lines at Pakistan's #1 e-commerce platform (~100M monthly sessions)",
                    "Delivered Alipay and local payment method integrations, expanding payment coverage by 40%",
                    "Reduced checkout abandonment by 22% through A/B testing programme and UX redesign",
                    "Managed Agile squads of 15 engineers and 3 designers using Scrum/Kanban frameworks",
                    "Collaborated with Alibaba product team in Hangzhou to localise global features for Pakistan market",
                ],
            },
            {
                "title": "Programme Manager",
                "company": "TapmadTV",
                "location": "Karachi, Pakistan",
                "period": "2016 – 2018",
                "bullets": [
                    "Led $3M digital transformation programme for Pakistan's first licensed OTT streaming platform",
                    "Built PMO governance from scratch: risk registers, RAID logs, milestone tracking, SteerCo reporting",
                    "Delivered platform launch on schedule across 5 technology workstreams (iOS, Android, web, CMS, CDN)",
                ],
            },
            {
                "title": "PMO Consultant",
                "company": "Various Clients (Dubizzle/OLX, Bank Alfalah, Zong/China Mobile)",
                "location": "UAE / Pakistan",
                "period": "2010 – 2016",
                "bullets": [
                    "Dubizzle (UAE): Led digital classified platform product development and programme management",
                    "Bank Alfalah: Digital banking transformation, core banking migration programme",
                    "Zong: Network expansion programme management for China Mobile's Pakistan subsidiary",
                ],
            },
        ]

        # Inject featured bullets from tailoring brief
        featured = {b["company"]: b["achievement"] for b in sections.get("featured_bullets", [])}
        new_bullets = sections.get("new_bullets", [])

        for exp in experiences:
            # Company/Title header
            p_title = doc.add_paragraph()
            title_run = p_title.add_run(exp["title"])
            title_run.bold = True
            title_run.font.size = Pt(10)
            p_title.space_after = Pt(0)

            p_company = doc.add_paragraph()
            co_run = p_company.add_run(f"{exp['company']} | {exp['location']} | {exp['period']}")
            co_run.italic = True
            co_run.font.size = Pt(9)
            co_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p_company.space_after = Pt(2)

            # Add featured bullet at top if relevant
            if exp["company"] in featured:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(featured[exp["company"]])
                run.font.size = Pt(9.5)
                run.bold = True  # Highlighted as tailored

            # Standard bullets
            for bullet in exp["bullets"][:5]:  # Limit per role
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(bullet)
                run.font.size = Pt(9.5)
                p.space_after = Pt(1)

            # Add new bullets for SimPaisa (from gap filling)
            if exp["company"] == "SimPaisa" and new_bullets:
                for nb in new_bullets[:3]:
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(nb.get("achievement", ""))
                    run.font.size = Pt(9.5)
                    run.bold = True

            doc.add_paragraph()

        # ── Certifications ──────────────────────────────────────────────────
        add_section_header("Certifications")
        cert_text = " | ".join(certs or ["PMP", "PMI-ACP", "CSPO", "CSM", "COBIT 5", "ITIL v3"])
        p_cert = doc.add_paragraph(cert_text)
        p_cert.style.font.size = Pt(10)
        p_cert.space_after = Pt(4)

        # ── Education ───────────────────────────────────────────────────────
        add_section_header("Education")
        edu_list = self.profile.get("education", [
            {"degree": "M.S. Applied Physics", "institution": "University of Karachi"},
            {"degree": "B.S. Physics, Statistics & Mathematics", "institution": "University of Karachi"},
        ])
        for edu in edu_list:
            p_edu = doc.add_paragraph(f"{edu['degree']} — {edu['institution']}")
            p_edu.style.font.size = Pt(10)
            p_edu.space_after = Pt(2)

        # Save
        settings = get_settings()
        os.makedirs(settings.output_resumes_dir, exist_ok=True)
        date_str = date.today().isoformat()
        safe_company = company.lower().replace(" ", "-").replace("/", "-")[:20]
        safe_role = job_title.lower().replace(" ", "-")[:20]
        filename = f"Rizwan_Zafar_{safe_company}_{safe_role}_{date_str}.docx"
        path = os.path.join(settings.output_resumes_dir, filename)
        doc.save(path)
        return path
