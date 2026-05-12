"""
api/workspace.py — FastAPI router for the Application Workspace surface.

The workspace is the page a user lands on when they click "Start
application process" on a /today action card. URL shape:

    /applications/{job_id}/workspace

This router serves the BUNDLE that page renders (one round-trip; the UI
does not chain calls), plus the mutation endpoints for the five tabs:

    GET    /workspace/{job_id}                  — full bundle
    POST   /workspace/{job_id}/build-resume     — enqueue G2
    POST   /workspace/{job_id}/edit-resume      — Quick tweak (Opus 4.7)
    POST   /workspace/{job_id}/save-resume-edit — persist user edit
    POST   /workspace/{job_id}/mark-applied     — applications row → applied
    GET    /workspace/{job_id}/resume.{format}  — md / pdf / docx download

All endpoints filter by `Depends(get_current_user)`. In single-user
mode this is Rizwan; in multi-tenant mode it's the Supabase JWT subject.

Design choices (recorded in api/WORKSPACE.md):

  • Bundle endpoint denormalises every piece of context the page needs,
    keyed off `job_id`. The cost is one extra DB round-trip vs. the
    smallest possible response — but the benefit is that the workspace
    page is a pure server-component fetch with no client-side waterfall.

  • Warm-intro paths reuse `agents.referral_graph.ReferralGraph.find_paths`
    via a fuzzy-resolved `target_company_id`. We don't 404 if there's no
    target_companies row yet — we degrade to `warm_intros_available: false`
    and the Network tab shows the import-CSV empty state.

  • Resume edit chat is stateless on the server today — the chat history
    rides in the request body. Persisting the conversation belongs to a
    later phase (add an `editor_chat_messages` table); see WORKSPACE.md
    for the contract.

  • PDF/DOCX downloads redirect to the existing `resume_pdf_url` /
    `resume_docx_url` URLs on `resume_builds` if those exist; otherwise
    501 with a clear message. PDF rendering from markdown is wired in
    Phase 3 alongside the cover-email rendering work.
"""
from __future__ import annotations

import logging
from datetime import date, datetime, timezone
from typing import Any, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import PlainTextResponse, RedirectResponse, Response
from pydantic import BaseModel, Field

from agents.referral_graph import (
    DEFAULT_MAX_HOPS,
    DEFAULT_MIN_STRENGTH,
    ReferralGraph,
    ReferralPath,
)
from agents.resume_edit_assistant import (
    CostCapExceeded,
    quick_tweak,
    rebuild_section as agent_rebuild_section,
)
from api.context import get_current_user
from api.queue import enqueue_g2_build
from api.users import User
from db.client import get_supabase

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/workspace", tags=["workspace"])


# ─── Request models ────────────────────────────────────────────────────────
class EditChatTurn(BaseModel):
    role: str = Field(description="user | assistant")
    content: str


class EditResumeBody(BaseModel):
    instruction: str = Field(min_length=1, max_length=2000)
    current_md: str = Field(min_length=1)
    chat_history: Optional[list[EditChatTurn]] = None
    mode: str = Field(default="quick_tweak", description="quick_tweak | rebuild_section | full_rebuild")


class SaveResumeEditBody(BaseModel):
    edited_md: str = Field(min_length=1)
    build_id: Optional[str] = Field(
        default=None,
        description="Resume build to save under. If omitted, the latest converged build for this job is used.",
    )


class RebuildSectionBody(BaseModel):
    section: str = Field(min_length=1, max_length=200)
    edit_intent: str = Field(min_length=1, max_length=2000)
    current_md: str = Field(min_length=1)


class FullRebuildBody(BaseModel):
    edit_intent: Optional[str] = Field(
        default=None,
        max_length=2000,
        description="Optional human instruction surfaced to the writer brief.",
    )
    max_cost_usd: Optional[float] = Field(
        default=None, ge=0.5, le=20.0,
        description="Override the per-build cost cap. Default uses settings.g2_max_cost_usd.",
    )


class MarkAppliedBody(BaseModel):
    applied_date: Optional[str] = Field(
        default=None,
        description="ISO date (YYYY-MM-DD). Defaults to today (UTC).",
    )
    notes: Optional[str] = None


# ─── Helpers ───────────────────────────────────────────────────────────────
def _get_job_for_user(db, *, job_id: int, user_id: UUID) -> dict[str, Any]:
    """Fetch a job row scoped to the current user. 404s on miss."""
    rows = (
        db.table("jobs")
        .select("*")
        .eq("id", job_id)
        .eq("user_id", str(user_id))
        .limit(1)
        .execute()
    ).data or []
    if not rows:
        raise HTTPException(status_code=404, detail="job_not_found")
    return rows[0]


def _get_application_for_job(db, *, job_id: int, user_id: UUID) -> Optional[dict[str, Any]]:
    """Most-recent application for this job. None if not yet applied."""
    rows = (
        db.table("applications")
        .select("*")
        .eq("job_id", job_id)
        .eq("user_id", str(user_id))
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    ).data or []
    return rows[0] if rows else None


def _get_latest_resume_build(db, *, job_id: int, user_id: UUID) -> Optional[dict[str, Any]]:
    """Latest build for this job + user. Prefer 'converged'; fall back to most recent."""
    rows = (
        db.table("resume_builds")
        .select(
            "id, job_id, application_id, company_name, persona_version, status, "
            "iterations, resume_md, resume_pdf_url, resume_docx_url, "
            "user_edited_md, user_edited_at, cost_usd_total, latency_ms_total, "
            "created_at, finalized_at, error"
        )
        .eq("job_id", job_id)
        .eq("user_id", str(user_id))
        .order("created_at", desc=True)
        .limit(10)
        .execute()
    ).data or []
    if not rows:
        return None
    converged = next((r for r in rows if r.get("status") == "converged"), None)
    return converged or rows[0]


def _get_persona_for_company(db, *, company_name: str, user_id: UUID) -> Optional[dict[str, Any]]:
    """Pull the company persona for the user. Tolerates missing rows."""
    if not company_name:
        return None
    rows = (
        db.table("company_personas")
        .select(
            "id, company_id, company_name, persona_version, "
            "ats_keyword_bank, success_patterns, failure_patterns, "
            "metadata, last_synthesized_at"
        )
        .eq("company_name", company_name)
        .eq("user_id", str(user_id))
        .order("persona_version", desc=True)
        .limit(1)
        .execute()
    ).data or []
    return rows[0] if rows else None


def _get_interview_prep_summary(db, *, job_id: int, user_id: UUID) -> Optional[dict[str, Any]]:
    """interview_prep summary if any pack exists for this job."""
    try:
        rows = (
            db.table("interview_prep")
            .select("id, application_id, status, prep_pack_url, round_type, round_number, created_at")
            .eq("job_id", job_id)
            .eq("user_id", str(user_id))
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        ).data or []
    except Exception as exc:
        # interview_prep may not have user_id wired in some tenants —
        # degrade gracefully.
        logger.warning("interview_prep lookup failed for job %s: %s", job_id, exc)
        return None
    if not rows:
        return None
    row = rows[0]
    has_pack = bool(row.get("prep_pack_url")) and row.get("status") == "converged"
    return {
        "has_pack": has_pack,
        "prep_pack_url": row.get("prep_pack_url"),
        "status": row.get("status"),
        "round_type": row.get("round_type"),
        "round_number": row.get("round_number"),
        "interview_prep_id": row.get("id"),
    }


def _resolve_target_company_id(db, *, company_name: str, user_id: UUID) -> Optional[str]:
    """Map a job's company name → target_companies.id for THIS user.

    Path-finder requires a target_company_id, not a name. We try exact,
    then case-insensitive, before giving up.

    Returns None if the user has no target_companies row for this
    company. The Network tab handles that case with an explicit "No warm
    intros to {company} yet" empty state.
    """
    if not company_name:
        return None
    try:
        # 1. Exact match on `company_name` column.
        rows = (
            db.table("target_companies")
            .select("id, company_name, name")
            .eq("user_id", str(user_id))
            .eq("company_name", company_name)
            .limit(1)
            .execute()
        ).data or []
        if rows:
            return str(rows[0]["id"])
        # 2. ILIKE fallback (handles "Stripe" vs "Stripe, Inc."-style drift).
        rows = (
            db.table("target_companies")
            .select("id, company_name, name")
            .eq("user_id", str(user_id))
            .ilike("company_name", f"%{company_name}%")
            .limit(1)
            .execute()
        ).data or []
        if rows:
            return str(rows[0]["id"])
        # 3. ILIKE on the display `name` column.
        rows = (
            db.table("target_companies")
            .select("id, company_name, name")
            .eq("user_id", str(user_id))
            .ilike("name", f"%{company_name}%")
            .limit(1)
            .execute()
        ).data or []
        if rows:
            return str(rows[0]["id"])
    except Exception as exc:
        logger.warning("target_company resolve failed for %s: %s", company_name, exc)
    return None


def _network_size_for_user(db, *, user_id: UUID) -> int:
    """Cheap HEAD-style count: how many people the user has on file.

    Used by the workspace bundle so the Network tab can render the
    "Import your LinkedIn CSV" empty state vs. "No warm intros to
    {company} yet — N contacts on file" message.
    """
    try:
        result = (
            db.table("people")
            .select("id", count="exact")
            .eq("user_id", str(user_id))
            .limit(1)
            .execute()
        )
        # Supabase-py returns the count on the response object.
        count = getattr(result, "count", None)
        if isinstance(count, int):
            return count
        return len(result.data or [])
    except Exception as exc:
        logger.warning("network size count failed: %s", exc)
        return 0


def _serialize_resume_build(row: Optional[dict[str, Any]]) -> Optional[dict[str, Any]]:
    """Trim the resume_build row to what the workspace UI actually needs."""
    if not row:
        return None
    return {
        "build_id": row.get("id"),
        "status": row.get("status"),
        "iterations": row.get("iterations"),
        "resume_md": row.get("resume_md"),
        "user_edited_md": row.get("user_edited_md"),
        "user_edited_at": row.get("user_edited_at"),
        "resume_pdf_url": row.get("resume_pdf_url"),
        "resume_docx_url": row.get("resume_docx_url"),
        "cost_usd_total": row.get("cost_usd_total"),
        "latency_ms_total": row.get("latency_ms_total"),
        "company_name": row.get("company_name"),
        "persona_version": row.get("persona_version"),
        "created_at": row.get("created_at"),
        "finalized_at": row.get("finalized_at"),
        "error": row.get("error"),
    }


def _serialize_persona(row: Optional[dict[str, Any]]) -> Optional[dict[str, Any]]:
    if not row:
        return None
    return {
        "company_id": row.get("company_id"),
        "company_name": row.get("company_name"),
        "persona_version": row.get("persona_version"),
        "ats_keyword_bank": row.get("ats_keyword_bank") or {},
        "success_patterns": row.get("success_patterns") or [],
        "failure_patterns": row.get("failure_patterns") or [],
        "metadata": row.get("metadata") or {},
        "last_synthesized_at": row.get("last_synthesized_at"),
    }


def _serialize_job(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": row.get("id"),
        "title": row.get("title"),
        "company": row.get("company"),
        "company_id": row.get("company_id"),
        "location": row.get("location"),
        "url": row.get("url"),
        "description": row.get("description"),
        "match_score": row.get("match_score"),
        "fit_details": row.get("fit_details") or {},
        "status": row.get("status"),
        "discovered_at": row.get("discovered_at"),
        "applied_at": row.get("applied_at"),
        "resume_generated_at": row.get("resume_generated_at"),
        "validation_status": row.get("validation_status"),
        "archetype": row.get("archetype"),
    }


# ─── GET /workspace/{job_id} ───────────────────────────────────────────────
@router.get("/{job_id}")
def get_workspace(
    job_id: int,
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Return the full workspace bundle for a job.

    One denormalised response so the page loads with no client-side
    waterfall. See module docstring for design rationale.
    """
    db = get_supabase()
    job = _get_job_for_user(db, job_id=job_id, user_id=user.id)
    application = _get_application_for_job(db, job_id=job_id, user_id=user.id)
    resume = _get_latest_resume_build(db, job_id=job_id, user_id=user.id)
    persona = _get_persona_for_company(db, company_name=job.get("company") or "", user_id=user.id)
    interview_prep = _get_interview_prep_summary(db, job_id=job_id, user_id=user.id)

    # Warm intros — only call the path-finder if we resolved a target.
    warm_paths: list[ReferralPath] = []
    target_company_id = _resolve_target_company_id(
        db, company_name=job.get("company") or "", user_id=user.id
    )
    if target_company_id:
        try:
            rg = ReferralGraph(user.id)
            warm_paths = rg.find_paths(
                target_company_id=target_company_id,
                max_hops=DEFAULT_MAX_HOPS,
                min_strength=DEFAULT_MIN_STRENGTH,
                limit=5,
            )
        except Exception as exc:
            logger.warning(
                "find_paths failed for job %s target %s: %s",
                job_id, target_company_id, exc,
            )
            warm_paths = []

    network_size = _network_size_for_user(db, user_id=user.id)

    return {
        "job": _serialize_job(job),
        "application": application,
        "resume": _serialize_resume_build(resume),
        "persona": _serialize_persona(persona),
        "interview_prep": interview_prep,
        "target_company_id": target_company_id,
        "warm_intro_paths": [p.model_dump() for p in warm_paths],
        "warm_intros_available": bool(warm_paths),
        "network_size": network_size,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


# ─── GET /workspace/{job_id}/comp-band ─────────────────────────────────────
#
# Phase 1.1 — compensation intelligence. Powers (a) the workspace Apply tab
# salary-expectation field, (b) G5 evaluation scoring (Phase 2), (c) G7
# application assistant salary handling (Phase 3), (d) G8 offer negotiation
# framing (Phase 4). One Perplexity Sonar call (~$0.02) per uncached
# (company, role, level, location) tuple per 30 days.

@router.get("/{job_id}/comp-band")
async def get_comp_band_for_job(
    job_id: int,
    location_override: Optional[str] = Query(
        default=None,
        description=(
            "Override the job's location (e.g. when the user is targeting a "
            "different geo than the listing itself — common for Remote-x "
            "ranges where one band per region exists)."
        ),
    ),
    level_override: Optional[str] = Query(
        default=None,
        description=(
            "Override the inferred level (Junior/Mid/Senior/Staff/Principal/"
            "Director). When omitted, the agent uses the job's `archetype` "
            "or falls back to extracting a level token from the title."
        ),
    ),
    user_target_total_comp: Optional[int] = Query(
        default=None,
        ge=10_000,
        le=10_000_000,
        description=(
            "User's own target total comp. Used to cap the strategy's anchor "
            "at the market p90 — preventing accidental ATS auto-filter for "
            "asking above-band."
        ),
    ),
    force_refresh: bool = Query(
        default=False,
        description=(
            "Bypass the 30-day cache and re-query Perplexity Sonar. "
            "Costs ~$0.02. Use sparingly."
        ),
    ),
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Comp-band for a job + the salary-field strategy.

    Bundle response so the dashboard renders the salary section in one
    round-trip:

        GET /workspace/123/comp-band?level_override=Senior

        {
          "band": {
            "company": "Stripe", "role": "Senior Product Manager",
            "level": "Senior", "location": "Dubai, UAE",
            "currency": "USD",
            "p25": 150000, "p50": 180000, "p75": 215000, "p90": 260000,
            "source_summary": "Glassdoor + Levels.fyi 2025-2026 ...",
            "source_citations": [{"url": "...", "title": "..."}],
            "cached": true, "age_days": 4, "cost_usd": 0,
            "confidence": "high", "cache_id": "..."
          },
          "strategy": {
            "approach": "range",
            "low": 195000, "high": 240000,
            "anchor": null,
            "framing": "I'm targeting total comp in the USD 195,000-240,000 range...",
            "rationale": "Market band is wide (44% spread above p50)..."
          },
          "generated_at": "2026-05-12T..."
        }

    The endpoint is scoped to the requesting user via
    `_get_job_for_user(...)` — RLS on `comp_cache` plus tenant filter on
    `jobs` together prevent cross-user leakage.
    """
    from agents.comp_research import (
        band_to_dict,
        get_comp_band,
        strategy_to_dict,
        suggest_salary_strategy,
    )

    db = get_supabase()
    job = _get_job_for_user(db, job_id=job_id, user_id=user.id)  # 404 if not theirs

    # Derive query inputs from the job row + overrides.
    company = job.get("company") or ""
    role = job.get("title") or ""
    location = location_override or job.get("location") or None
    # Level: prefer override, else heuristic from archetype/title.
    level = level_override
    if not level:
        archetype = job.get("archetype")
        if isinstance(archetype, str) and archetype:
            level = archetype
        else:
            # Crude title-based fallback. Real Phase-2 level extraction will
            # use a structured archetype detector — for now, this catches the
            # 80% case (senior / staff / principal / lead / head / director).
            title_lc = (role or "").lower()
            for token in ("principal", "staff", "head of", "director",
                          "lead", "senior", "junior", "intern"):
                if token in title_lc:
                    level = token.title()
                    break

    if not company or not role:
        raise HTTPException(
            status_code=409,
            detail={
                "code": "missing_job_fields",
                "message": (
                    f"Cannot research comp without a company and a role title. "
                    f"job={job_id}: company={company!r} title={role!r}."
                ),
            },
        )

    band = await get_comp_band(
        company=company,
        role=role,
        level=level,
        location=location,
        user_id=user.id,
        force_refresh=force_refresh,
    )
    strategy = suggest_salary_strategy(band, user_target=user_target_total_comp)

    return {
        "band": band_to_dict(band),
        "strategy": strategy_to_dict(strategy),
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


# ─── POST /workspace/{job_id}/build-resume ─────────────────────────────────
@router.post("/{job_id}/build-resume")
def build_resume(
    job_id: int,
    force: bool = Query(default=False),
    max_cost_usd: Optional[float] = Query(default=None, ge=0.5, le=20.0),
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Enqueue a fresh G2 build for this job.

    Idempotent on (user_id, kind=g2_resume, payload) — the queue's
    `_enqueue_or_dedup` guard returns the existing run id if one is
    already queued/running.

    Returns the jobs_runs.id; the UI polls `/jobs-runs/{run_id}` (already
    shipped in api/jobs_runs.py) every ~8s until status is terminal.
    """
    db = get_supabase()
    job = _get_job_for_user(db, job_id=job_id, user_id=user.id)  # 404 if not theirs

    # 2026-05-12: pre-flight staleness guardrail.
    #
    # G2 costs ~$0.50-$1.50 in LLM tokens per build (5 models, 3 iterations,
    # ATS critics A/B, persona critic, polisher). Building a resume against
    # a job posting that is closed, expired, or failed validation is
    # **direct financial waste** — the user cannot apply to it anyway.
    #
    # Production incident (2026-05-12): OKX job 1641 "Senior Product Manager,
    # Payment" was scraped from LinkedIn 1 year after its original posting
    # date. The freshness pipeline didn't catch it because LinkedIn renders
    # post age as plain text ("1 year ago") not as metadata. The user built
    # a resume against it before noticing — pure waste of LLM spend.
    #
    # The validation_failed='stale_per_description' guard now closes such
    # listings at discovery time (see agents/job_validation.py
    # _detect_stale_age_marker), but jobs already in the DB before that fix
    # remain stuck open. This is the layer that protects the user's wallet.
    #
    # Override: pass force=true to build against a closed posting (e.g. the
    # user has insider knowledge that the role is actually still active).
    if not force:
        if job.get("posting_closed_at") is not None:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "posting_closed",
                    "message": (
                        "This job posting is marked closed"
                        + (
                            f" ({job['validation_failed']})"
                            if job.get("validation_failed")
                            else ""
                        )
                        + ". Building a resume costs ~$1 in LLM tokens and "
                        "you cannot apply to a closed posting. Pass "
                        "force=true to override."
                    ),
                    "posting_closed_at": job.get("posting_closed_at"),
                    "validation_failed": job.get("validation_failed"),
                },
            )
        if job.get("validation_failed"):
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "validation_failed",
                    "message": (
                        f"This job failed validation "
                        f"({job['validation_failed']}). Building a resume is "
                        "likely wasted spend. Pass force=true to override."
                    ),
                    "validation_failed": job["validation_failed"],
                },
            )

    try:
        run_id = enqueue_g2_build(
            user_id=user.id,
            job_id=job_id,
            force=force,
            max_cost_usd=max_cost_usd,
        )
    except RuntimeError as exc:
        # Redis/RQ not installed in this environment — surface clearly.
        raise HTTPException(
            status_code=503,
            detail=f"queue_unavailable: {exc}",
        ) from exc
    return {
        "run_id": run_id,
        "status": "queued",
        "kind": "g2_resume",
        "job_id": job_id,
        "force": force,
        "max_cost_usd": max_cost_usd,
        "poll_url": f"/jobs-runs/{run_id}",
    }


# ─── POST /workspace/{job_id}/edit-resume ──────────────────────────────────
@router.post("/{job_id}/edit-resume")
async def edit_resume(
    job_id: int,
    body: EditResumeBody,
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Apply ONE edit instruction to the resume markdown.

    Phase 2 ships only the `quick_tweak` mode. The other two modes return
    501 with a clear "coming next session" message — keeps the chat panel
    UI stable while the backend catches up.

    Stateless on the server: the chat history rides in the request body.
    Persistence to a future `editor_chat_messages` table is documented
    in api/WORKSPACE.md — out of scope for Phase 2.
    """
    db = get_supabase()
    job = _get_job_for_user(db, job_id=job_id, user_id=user.id)

    # Rebuild-section / full-rebuild now have their own endpoints (the
    # chat panel routes there directly). Keep these mode codes accepted
    # in /edit-resume so old clients get a clear redirect, not a 501.
    if body.mode == "rebuild_section":
        raise HTTPException(
            status_code=400,
            detail={
                "code": "use_dedicated_endpoint",
                "mode": "rebuild_section",
                "message": (
                    "Use POST /workspace/{job_id}/rebuild-section with "
                    "{section, edit_intent, current_md} — the chat-shape "
                    "endpoint no longer dispatches rebuilds."
                ),
            },
        )
    if body.mode == "full_rebuild":
        raise HTTPException(
            status_code=400,
            detail={
                "code": "use_dedicated_endpoint",
                "mode": "full_rebuild",
                "message": (
                    "Use POST /workspace/{job_id}/full-rebuild with "
                    "{edit_intent?} — the chat-shape endpoint no longer "
                    "dispatches rebuilds."
                ),
            },
        )
    if body.mode != "quick_tweak":
        raise HTTPException(status_code=400, detail=f"unknown_mode:{body.mode}")

    persona = _get_persona_for_company(db, company_name=job.get("company") or "", user_id=user.id)
    chat_history = (
        [{"role": t.role, "content": t.content} for t in body.chat_history]
        if body.chat_history else None
    )

    try:
        result = await quick_tweak(
            current_md=body.current_md,
            instruction=body.instruction,
            persona=persona,
            jd={
                "company": job.get("company"),
                "title": job.get("title"),
                "description": job.get("description"),
            },
            chat_history=chat_history,
        )
    except CostCapExceeded as exc:
        # 429 is the right shape: "you tried, the system refused, try again".
        raise HTTPException(
            status_code=429,
            detail={
                "code": "cost_cap_exceeded",
                "message": str(exc),
            },
        ) from exc
    except ValueError as exc:
        # Empty input or malformed model output.
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("quick_tweak failed for job %s", job_id)
        raise HTTPException(
            status_code=502,
            detail=f"resume_editor_unavailable: {type(exc).__name__}",
        ) from exc

    return result


# ─── POST /workspace/{job_id}/rebuild-section ──────────────────────────────
@router.post("/{job_id}/rebuild-section")
async def rebuild_section_endpoint(
    job_id: int,
    body: RebuildSectionBody,
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Rebuild ONE H2 section of the resume synchronously.

    Runs the writer → critic → polish mini-graph in
    `agents.resume_edit_assistant.rebuild_section`. Returns the FULL
    updated markdown (the rest of the resume is preserved verbatim by
    markdown-splice).

    Synchronous on purpose: the operation is bounded at 60s, so a long
    HTTP request is cleaner than enqueue + poll for a flow that always
    completes within a single tab session. If we ever push this past
    60s wall-clock we'll switch to the queue path (warm_start_md +
    rebuild_scope='section' on enqueue_g2_build is already plumbed).

    On timeout: 504-shaped error so the UI can suggest "switch to a
    smaller section or fall back to Quick tweak".
    """
    db = get_supabase()
    job = _get_job_for_user(db, job_id=job_id, user_id=user.id)

    persona = _get_persona_for_company(db, company_name=job.get("company") or "", user_id=user.id)

    try:
        result = await agent_rebuild_section(
            current_md=body.current_md,
            section=body.section,
            edit_intent=body.edit_intent,
            persona=persona,
            jd={
                "company": job.get("company"),
                "title": job.get("title"),
                "description": job.get("description"),
            },
        )
    except CostCapExceeded as exc:
        raise HTTPException(
            status_code=429,
            detail={
                "code": "cost_cap_exceeded",
                "message": str(exc),
            },
        ) from exc
    except TimeoutError as exc:
        # 504-ish: the gateway side timed out before producing the result.
        raise HTTPException(
            status_code=504,
            detail={
                "code": "rebuild_section_timeout",
                "message": str(exc),
            },
        ) from exc
    except ValueError as exc:
        # Empty input, missing section, or malformed model output.
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("rebuild_section failed for job %s section=%s", job_id, body.section)
        raise HTTPException(
            status_code=502,
            detail=f"rebuild_section_unavailable: {type(exc).__name__}",
        ) from exc

    return result


# ─── POST /workspace/{job_id}/full-rebuild ─────────────────────────────────
@router.post("/{job_id}/full-rebuild")
def full_rebuild_endpoint(
    job_id: int,
    body: FullRebuildBody,
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Enqueue a full G2 rebuild from scratch.

    Thin wrapper around `enqueue_g2_build(force=True)` so the dedup hash
    differs from any prior run. The optional `edit_intent` is recorded
    in the payload and surfaced to the writer node's brief — useful when
    the user has a north-star instruction ("make this more product-led")
    but still wants the full ensemble pipeline.

    Returns the jobs_runs id; the UI polls /jobs-runs/{run_id} every ~8s
    via the same flow as the "Build resume" button on the Resume tab.
    """
    db = get_supabase()
    _get_job_for_user(db, job_id=job_id, user_id=user.id)  # 404 if not theirs

    try:
        run_id = enqueue_g2_build(
            user_id=user.id,
            job_id=job_id,
            force=True,
            max_cost_usd=body.max_cost_usd,
            edit_intent=body.edit_intent,
            rebuild_scope="full",
        )
    except RuntimeError as exc:
        raise HTTPException(
            status_code=503,
            detail=f"queue_unavailable: {exc}",
        ) from exc

    return {
        "run_id": run_id,
        "status": "queued",
        "kind": "g2_resume",
        "job_id": job_id,
        "force": True,
        "rebuild_scope": "full",
        "edit_intent": body.edit_intent,
        "max_cost_usd": body.max_cost_usd,
        "poll_url": f"/jobs-runs/{run_id}",
    }


# ─── POST /workspace/{job_id}/save-resume-edit ─────────────────────────────
@router.post("/{job_id}/save-resume-edit")
def save_resume_edit(
    job_id: int,
    body: SaveResumeEditBody,
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Persist the user-edited markdown onto the resume_builds row.

    This is the EXACT same write the existing
    PATCH /resume-builds/{id}/markdown does, but scoped to the workspace
    (job_id is the source of truth on this surface; the UI doesn't care
    about build ids unless explicitly switching builds).

    If `build_id` is supplied we honour it — otherwise we save against
    the latest converged build for this job + user.
    """
    db = get_supabase()
    job = _get_job_for_user(db, job_id=job_id, user_id=user.id)

    build_row = None
    if body.build_id:
        rows = (
            db.table("resume_builds")
            .select("*")
            .eq("id", body.build_id)
            .eq("user_id", str(user.id))
            .eq("job_id", job_id)
            .limit(1)
            .execute()
        ).data or []
        if rows:
            build_row = rows[0]
    if build_row is None:
        build_row = _get_latest_resume_build(db, job_id=job_id, user_id=user.id)
    if build_row is None:
        raise HTTPException(status_code=404, detail="no_resume_build_for_job")

    update = {
        "user_edited_md": body.edited_md,
        "user_edited_at": datetime.now(timezone.utc).isoformat(),
    }
    result = (
        db.table("resume_builds")
        .update(update)
        .eq("id", build_row["id"])
        .execute()
    )
    saved = (result.data or [None])[0]
    if not saved:
        # Some postgrest versions don't echo on update; re-read.
        saved_rows = (
            db.table("resume_builds")
            .select("*")
            .eq("id", build_row["id"])
            .limit(1)
            .execute()
        ).data or []
        saved = saved_rows[0] if saved_rows else build_row

    return {
        "saved": True,
        "build_id": build_row["id"],
        "job_id": job_id,
        "company": job.get("company"),
        "user_edited_at": saved.get("user_edited_at"),
        "byte_size": len((body.edited_md or "").encode("utf-8")),
    }


# ─── POST /workspace/{job_id}/mark-applied ─────────────────────────────────
@router.post("/{job_id}/mark-applied")
def mark_applied(
    job_id: int,
    body: MarkAppliedBody,
    user: User = Depends(get_current_user),
) -> dict[str, Any]:
    """Move the application to 'applied'. Creates the row if absent.

    After this call the home action card (resume_ready) drops off /today
    and the Applications board shows the row in its 'applied' column.
    """
    db = get_supabase()
    job = _get_job_for_user(db, job_id=job_id, user_id=user.id)
    applied_iso = body.applied_date or date.today().isoformat()

    existing = _get_application_for_job(db, job_id=job_id, user_id=user.id)
    if existing:
        update: dict[str, Any] = {
            "status": "applied",
            "applied_date": applied_iso,
        }
        if body.notes is not None:
            update["notes"] = body.notes
        result = (
            db.table("applications")
            .update(update)
            .eq("id", existing["id"])
            .execute()
        )
        row = (result.data or [None])[0] or {**existing, **update}
        # Update the jobs row's applied_at + status so /today drops the card.
        try:
            db.table("jobs").update({
                "status": "applied",
                "applied_at": datetime.now(timezone.utc).isoformat(),
            }).eq("id", job_id).eq("user_id", str(user.id)).execute()
        except Exception as exc:
            logger.warning("jobs.applied_at backfill failed: %s", exc)
        return {"created": False, "updated": True, "application": row}

    # Create new
    new_row: dict[str, Any] = {
        "user_id": str(user.id),
        "job_id": job_id,
        "company": job.get("company") or "",
        "role": job.get("title") or "",
        "status": "applied",
        "applied_date": applied_iso,
        "score": (job.get("match_score") or 0) / 20.0,  # 0-100 → 0-5
        "company_id": job.get("company_id"),
    }
    if body.notes is not None:
        new_row["notes"] = body.notes

    result = db.table("applications").insert(new_row).execute()
    row = (result.data or [None])[0] or new_row

    # Mirror onto the jobs row — same as the existing /applications POST.
    try:
        db.table("jobs").update({
            "status": "applied",
            "applied_at": datetime.now(timezone.utc).isoformat(),
        }).eq("id", job_id).eq("user_id", str(user.id)).execute()
    except Exception as exc:
        logger.warning("jobs.applied_at backfill failed: %s", exc)

    return {"created": True, "updated": False, "application": row}


# ─── GET /workspace/{job_id}/resume.{format} ───────────────────────────────
@router.get("/{job_id}/resume.{fmt}")
def download_resume(
    job_id: int,
    fmt: str,
    user: User = Depends(get_current_user),
):
    """Server-side download proxy for resume artifacts.

    Phase 2 contract:
      • md   → returns the latest user_edited_md (or resume_md) inline
               as text/markdown. Always works if a build exists.
      • pdf  → 302 redirect to resume_pdf_url if present, else 501.
      • docx → 302 redirect to resume_docx_url if present, else 501.

    Phase 3 will add server-side rendering (markdown → PDF/DOCX) so that
    user edits flow through to PDF without re-running G2.
    """
    fmt_norm = (fmt or "").strip().lower()
    if fmt_norm not in ("md", "pdf", "docx"):
        raise HTTPException(status_code=400, detail=f"unsupported_format:{fmt_norm}")

    db = get_supabase()
    _get_job_for_user(db, job_id=job_id, user_id=user.id)
    build = _get_latest_resume_build(db, job_id=job_id, user_id=user.id)
    if not build:
        raise HTTPException(status_code=404, detail="no_resume_build_for_job")

    if fmt_norm == "md":
        md = build.get("user_edited_md") or build.get("resume_md") or ""
        if not md:
            raise HTTPException(status_code=404, detail="resume_markdown_empty")
        filename = f"resume-{job_id}.md"
        return PlainTextResponse(
            md,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if fmt_norm == "pdf":
        # First preference: a pre-rendered PDF in storage.
        url = build.get("resume_pdf_url")
        if url:
            return RedirectResponse(url=url, status_code=302)

        # 2026-05-12: on-demand render. Until today this branch returned 501,
        # which broke the "Download PDF" button for every build that didn't
        # emit a pre-rendered PDF (which is most of them — G2 export hardcodes
        # resume_pdf_url=None at g2_nodes.py:1148 because pandoc/LaTeX isn't
        # installed in the slim Dockerfile). We now render markdown → PDF
        # synchronously and stream the bytes back. This means user edits
        # (user_edited_md) flow through to PDF without re-running G2.
        md_source = build.get("user_edited_md") or build.get("resume_md") or ""
        if not md_source:
            raise HTTPException(status_code=404, detail="resume_markdown_empty")
        try:
            pdf_bytes = _render_resume_md_to_pdf(md_source)
        except Exception as e:
            logger.exception("on-demand PDF render failed for job_id=%s", job_id)
            raise HTTPException(
                status_code=502,
                detail={
                    "code": "pdf_render_failed",
                    "message": f"server-side PDF render failed: {type(e).__name__}",
                },
            )
        filename = f"resume-{job_id}.pdf"
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # docx — same on-demand pattern as PDF above.
    url = build.get("resume_docx_url")
    if url:
        return RedirectResponse(url=url, status_code=302)
    md_source = build.get("user_edited_md") or build.get("resume_md") or ""
    if not md_source:
        raise HTTPException(status_code=404, detail="resume_markdown_empty")
    try:
        docx_bytes = _render_resume_md_to_docx(md_source)
    except Exception as e:
        logger.exception("on-demand DOCX render failed for job_id=%s", job_id)
        raise HTTPException(
            status_code=502,
            detail={
                "code": "docx_render_failed",
                "message": f"server-side DOCX render failed: {type(e).__name__}",
            },
        )
    filename = f"resume-{job_id}.docx"
    return Response(
        content=docx_bytes,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Markdown → PDF / DOCX render helpers ──────────────────────────────────
#
# Both renderers are pure-Python (no pandoc, no LaTeX, no system libs beyond
# what python-docx + xhtml2pdf + Pillow already bring in transitively). They
# work on the slim Dockerfile we deploy to Railway.
#
# Design notes:
#   * Resumes are small (1-3 pages, mostly headings + bullet lists). We don't
#     need pixel-perfect output — we need a clean, single-column document that
#     a recruiter can actually open. xhtml2pdf + python-docx hit that bar.
#   * On-demand rendering means user edits to `user_edited_md` flow through
#     to the downloaded artifact without re-running G2. That's the right UX
#     for "tweak a bullet → download → send" iteration.

def _render_resume_md_to_pdf(md: str) -> bytes:
    """Markdown → PDF bytes via reportlab platypus (pure-Python pipeline).

    Parses markdown structure (headings, bullet lists, emphasis) and emits
    flowables onto an A4 page. Deliberately keeps the styling restrained —
    Helvetica, single column, mild colour palette — so the output reads as
    a professional resume rather than a styled blog post. Caller wraps
    Exception → 502 so render failures surface cleanly.
    """
    import io
    import re
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        ListFlowable, ListItem,
    )

    base = getSampleStyleSheet()
    INK = HexColor("#0f1928")
    BODY = HexColor("#1d2433")
    RULE = HexColor("#cfd6e0")

    style_h1 = ParagraphStyle(
        "RH1", parent=base["Title"], fontName="Helvetica-Bold",
        fontSize=20, leading=24, textColor=INK,
        spaceBefore=0, spaceAfter=4, alignment=TA_LEFT,
    )
    style_h2 = ParagraphStyle(
        "RH2", parent=base["Heading2"], fontName="Helvetica-Bold",
        fontSize=11.5, leading=14, textColor=INK,
        spaceBefore=10, spaceAfter=4,
    )
    style_h3 = ParagraphStyle(
        "RH3", parent=base["Heading3"], fontName="Helvetica-Bold",
        fontSize=10.5, leading=13, textColor=INK,
        spaceBefore=6, spaceAfter=2,
    )
    style_body = ParagraphStyle(
        "RBody", parent=base["BodyText"], fontName="Helvetica",
        fontSize=10, leading=13.5, textColor=BODY,
        spaceBefore=2, spaceAfter=2,
    )
    style_bullet = ParagraphStyle(
        "RBullet", parent=style_body, leftIndent=0, bulletIndent=0,
        spaceBefore=1, spaceAfter=1,
    )

    def _inline(text: str) -> str:
        """Convert markdown **bold** / *italic* to reportlab mini-HTML.

        Reportlab's Paragraph supports <b>, <i>, <font color> tags. We
        escape `&`/`<`/`>` first so a job description containing "<5%"
        doesn't blow up the parser.
        """
        t = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # **bold**
        t = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", t)
        # *italic* (after bold so ** wasn't a single *)
        t = re.sub(r"(?<![*\w])\*([^*\n]+?)\*(?!\w)", r"<i>\1</i>", t)
        # Strip stray markdown link syntax [text](url) → text (url)
        t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", t)
        return t

    flow: list = []
    buf_lines: list[str] = []

    def _flush_paragraph() -> None:
        if buf_lines:
            joined = " ".join(line.strip() for line in buf_lines if line.strip())
            if joined:
                flow.append(Paragraph(_inline(joined), style_body))
            buf_lines.clear()

    def _flush_list(items: list[str]) -> None:
        if not items:
            return
        bullets = [
            ListItem(
                Paragraph(_inline(i), style_bullet),
                leftIndent=10, bulletColor=INK,
            )
            for i in items
        ]
        flow.append(
            ListFlowable(
                bullets,
                bulletType="bullet",
                bulletFontName="Helvetica",
                start="•",
                leftIndent=14,
                bulletFontSize=9,
                spaceBefore=2, spaceAfter=6,
            )
        )

    pending_list: list[str] = []
    for raw in (md or "").splitlines():
        line = raw.rstrip()
        bullet_match = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if bullet_match:
            _flush_paragraph()
            pending_list.append(bullet_match.group(1))
            continue
        # Anything that's not a bullet flushes the pending list.
        if pending_list:
            _flush_list(pending_list)
            pending_list = []
        if not line.strip():
            _flush_paragraph()
            continue
        if line.startswith("### "):
            _flush_paragraph()
            flow.append(Paragraph(_inline(line[4:].strip()), style_h3))
        elif line.startswith("## "):
            _flush_paragraph()
            flow.append(Paragraph(_inline(line[3:].strip()), style_h2))
            flow.append(HRFlowable(width="100%", thickness=0.5, color=RULE,
                                   spaceBefore=1, spaceAfter=4))
        elif line.startswith("# "):
            _flush_paragraph()
            flow.append(Paragraph(_inline(line[2:].strip()), style_h1))
        elif line.strip() == "---":
            _flush_paragraph()
            flow.append(HRFlowable(width="100%", thickness=0.5, color=RULE,
                                   spaceBefore=4, spaceAfter=4))
        else:
            buf_lines.append(line)

    _flush_list(pending_list)
    _flush_paragraph()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title="Resume", author="Job Hunt",
    )
    doc.build(flow)
    return buf.getvalue()


def _render_resume_md_to_docx(md: str) -> bytes:
    """Markdown → DOCX bytes via python-docx (already in requirements.txt).

    Lightweight markdown parsing: headings, bullets, emphasis. Not pandoc-
    fidelity but produces a clean editable Word document with section breaks
    a recruiter can paste straight into their ATS.
    """
    import io
    import re
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Helvetica"
    style_normal.font.size = Pt(10.5)

    def _add_inline(paragraph, text: str) -> None:
        # Minimal inline-emphasis parse: **bold** and *italic*.
        # Pattern: split on **bold** then within each chunk on *italic*.
        for chunk in re.split(r"(\*\*[^*]+\*\*)", text):
            if not chunk:
                continue
            if chunk.startswith("**") and chunk.endswith("**"):
                r = paragraph.add_run(chunk[2:-2])
                r.bold = True
                continue
            for sub in re.split(r"(\*[^*]+\*)", chunk):
                if not sub:
                    continue
                if sub.startswith("*") and sub.endswith("*"):
                    r = paragraph.add_run(sub[1:-1])
                    r.italic = True
                else:
                    paragraph.add_run(sub)

    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(); p.style = doc.styles["Heading 3"]
            _add_inline(p, line[4:].strip())
        elif line.startswith("## "):
            p = doc.add_paragraph(); p.style = doc.styles["Heading 2"]
            _add_inline(p, line[3:].strip())
        elif line.startswith("# "):
            p = doc.add_paragraph(); p.style = doc.styles["Heading 1"]
            _add_inline(p, line[2:].strip())
        elif re.match(r"^\s*[-*+]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, re.sub(r"^\s*[-*+]\s+", "", line))
        elif line.strip() == "---":
            # crude separator
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
