"""
resume_agents/render.py — pure-Python Markdown → DOCX / PDF renderers.

Why this exists
---------------
G2 (`export_node`) historically shelled out to `pandoc` to turn the
resume markdown into a DOCX. pandoc is NOT installed in either the API
(`Dockerfile`) or worker (`Dockerfile.worker`) image, so the subprocess
raised, `_local_md_to_docx` returned `b""`, and every build finalized
with `resume_docx_url=None`. The resume *text* lived in
`resume_builds.resume_md` all along, but there was no downloadable
artifact and no on-demand renderer — so a "converged" build looked
empty to the user. (Root-caused 2026-05-31.)

This module renders straight from the markdown string using libraries
that are ALREADY pinned in requirements.txt:

  - DOCX via `python-docx`  (requirements.txt: python-docx>=1.2.0)
  - PDF  via `reportlab`    (requirements.txt: reportlab>=4.0.0)
                            + `markdown`  (requirements.txt: markdown>=3.5)

No system binaries, no network, no Supabase Storage round-trip. Both
functions are defensive: they degrade gracefully (return b"" + log)
rather than raise, so a render failure can never crash a build or a
download request.

Markdown subset supported (matches what the G2 polisher emits — see a
real sample in resume_builds.resume_md):

    # H1                     → document title
    ## H2                    → section heading
    ### H3                   → sub-heading (job/company line)
    *italic*  **bold**       → inline runs
    - bullet / * bullet      → bullet list item
    plain paragraph          → body text
    ---                      → horizontal rule (DOCX: spacer; PDF: rule)
    | a | b |                → rendered as plain text line (resumes use
                               pipes as visual separators, not tables)

The goal is recruiter-shippable, ATS-friendly output — clean single
column, standard fonts — not pixel-perfect fidelity.
"""
from __future__ import annotations

import logging
import re
from io import BytesIO

logger = logging.getLogger(__name__)

__all__ = ["markdown_to_docx_bytes", "markdown_to_pdf_bytes"]


# ─────────────────────────────────────────────────────────────────────
# Shared lightweight inline-markdown parsing
# ─────────────────────────────────────────────────────────────────────
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _strip_inline(text: str) -> str:
    """Drop markdown emphasis markers for contexts that can't style
    inline (e.g. PDF flowable text we pass through Paragraph markup
    separately). Kept simple on purpose."""
    text = _BOLD_RE.sub(r"\1", text)
    text = _ITALIC_RE.sub(r"\1", text)
    return text


def _iter_bold_italic_runs(text: str):
    """Yield (substring, bold, italic) tuples for a single line.

    Handles **bold** and *italic* (non-overlapping, bold checked first).
    Anything not matched is emitted as a plain run. Used by the DOCX
    renderer which styles per-run.
    """
    # Tokenize on bold first, then italic within the non-bold spans.
    idx = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > idx:
            yield from _emit_italic(text[idx:m.start()])
        yield (m.group(1), True, False)
        idx = m.end()
    if idx < len(text):
        yield from _emit_italic(text[idx:])


def _emit_italic(text: str):
    idx = 0
    for m in _ITALIC_RE.finditer(text):
        if m.start() > idx:
            yield (text[idx:m.start()], False, False)
        yield (m.group(1), False, True)
        idx = m.end()
    if idx < len(text):
        yield (text[idx:], False, False)


def _is_bullet(line: str) -> bool:
    s = line.lstrip()
    return s.startswith("- ") or s.startswith("* ")


def _bullet_text(line: str) -> str:
    return line.lstrip()[2:].strip()


# ─────────────────────────────────────────────────────────────────────
# DOCX renderer (python-docx)
# ─────────────────────────────────────────────────────────────────────
def markdown_to_docx_bytes(md: str) -> bytes:
    """Render resume markdown to a .docx byte string. Returns b"" on
    failure (never raises)."""
    if not md or not md.strip():
        return b""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:  # pragma: no cover - python-docx is pinned
        logger.warning(f"python-docx unavailable for DOCX render: {e}")
        return b""

    try:
        doc = Document()

        # Tighten default style to a clean, ATS-safe single column.
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        for raw in md.splitlines():
            line = raw.rstrip()
            stripped = line.strip()

            if not stripped:
                continue

            # Horizontal rule → a thin spacer paragraph.
            if re.fullmatch(r"-{3,}", stripped) or re.fullmatch(r"\*{3,}", stripped):
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(2)
                continue

            # Headings
            if stripped.startswith("### "):
                p = doc.add_heading(level=3)
                _add_runs(p, stripped[4:])
                continue
            if stripped.startswith("## "):
                p = doc.add_heading(level=2)
                _add_runs(p, stripped[3:])
                continue
            if stripped.startswith("# "):
                # Title — center it, larger.
                p = doc.add_heading(level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_runs(p, stripped[2:])
                continue

            # Bullets
            if _is_bullet(stripped):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, _bullet_text(stripped))
                continue

            # Plain paragraph (includes pipe-separated "contact" lines)
            p = doc.add_paragraph()
            _add_runs(p, stripped)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.warning(f"DOCX render failed: {e}")
        return b""


def _add_runs(paragraph, text: str) -> None:
    """Add bold/italic-aware runs to a python-docx paragraph."""
    for substr, bold, italic in _iter_bold_italic_runs(text):
        if not substr:
            continue
        run = paragraph.add_run(substr)
        run.bold = bool(bold)
        run.italic = bool(italic)


# ─────────────────────────────────────────────────────────────────────
# PDF renderer (reportlab + markdown)
# ─────────────────────────────────────────────────────────────────────
def markdown_to_pdf_bytes(md: str) -> bytes:
    """Render resume markdown to a PDF byte string. Returns b"" on
    failure (never raises).

    Uses reportlab's platypus Flowables. Inline emphasis is converted to
    reportlab's mini-HTML (<b>/<i>) markup, which Paragraph understands.
    """
    if not md or not md.strip():
        return b""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem,
        )
    except Exception as e:  # pragma: no cover - reportlab is pinned
        logger.warning(f"reportlab unavailable for PDF render: {e}")
        return b""

    try:
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle(
            "ResumeH1", parent=styles["Title"], fontSize=18, leading=22,
            alignment=TA_CENTER, spaceAfter=4,
        )
        h2 = ParagraphStyle(
            "ResumeH2", parent=styles["Heading2"], fontSize=12, leading=15,
            textColor=HexColor("#1a1a1a"), spaceBefore=8, spaceAfter=3,
        )
        h3 = ParagraphStyle(
            "ResumeH3", parent=styles["Heading3"], fontSize=11, leading=14,
            spaceBefore=5, spaceAfter=2,
        )
        body = ParagraphStyle(
            "ResumeBody", parent=styles["Normal"], fontSize=9.5, leading=13,
            spaceAfter=2,
        )

        story = []
        bullets: list = []

        def flush_bullets():
            nonlocal bullets
            if bullets:
                story.append(ListFlowable(
                    [ListItem(Paragraph(_md_inline_to_rl(b), body), leftIndent=10)
                     for b in bullets],
                    bulletType="bullet", start="•", leftIndent=12,
                ))
                bullets = []

        for raw in md.splitlines():
            stripped = raw.strip()
            if not stripped:
                continue

            if re.fullmatch(r"-{3,}", stripped) or re.fullmatch(r"\*{3,}", stripped):
                flush_bullets()
                story.append(HRFlowable(
                    width="100%", thickness=0.5, color=HexColor("#cccccc"),
                    spaceBefore=3, spaceAfter=3,
                ))
                continue

            if stripped.startswith("### "):
                flush_bullets()
                story.append(Paragraph(_md_inline_to_rl(stripped[4:]), h3))
                continue
            if stripped.startswith("## "):
                flush_bullets()
                story.append(Paragraph(_md_inline_to_rl(stripped[3:]), h2))
                continue
            if stripped.startswith("# "):
                flush_bullets()
                story.append(Paragraph(_md_inline_to_rl(stripped[2:]), h1))
                continue

            if _is_bullet(stripped):
                bullets.append(_bullet_text(stripped))
                continue

            flush_bullets()
            story.append(Paragraph(_md_inline_to_rl(stripped), body))

        flush_bullets()

        buf = BytesIO()
        SimpleDocTemplate(
            buf, pagesize=letter,
            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
            topMargin=0.6 * inch, bottomMargin=0.6 * inch,
            title="Resume",
        ).build(story)
        return buf.getvalue()
    except Exception as e:
        logger.warning(f"PDF render failed: {e}")
        return b""


def _md_inline_to_rl(text: str) -> str:
    """Convert **bold**/*italic* to reportlab Paragraph markup (<b>/<i>),
    escaping XML-significant chars first so '&', '<', '>' in the resume
    text don't break the mini-HTML parser."""
    text = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    text = _BOLD_RE.sub(r"<b>\1</b>", text)
    text = _ITALIC_RE.sub(r"<i>\1</i>", text)
    return text
